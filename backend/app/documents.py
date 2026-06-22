"""Extract plain text from uploaded CV / JD files (PDF, DOCX, TXT)."""
from __future__ import annotations

import io
import logging

from app.config import get_settings

logger = logging.getLogger("recruitiq.documents")
settings = get_settings()

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}


def validate_file(filename: str, content: bytes) -> None:
    ext = _ext(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext or filename}'. Allowed: PDF, DOCX, TXT."
        )
    max_bytes = settings.max_upload_size_mb * 1024 * 1024
    if len(content) > max_bytes:
        raise ValueError(
            f"'{filename}' is {len(content) // (1024 * 1024)}MB, exceeds "
            f"{settings.max_upload_size_mb}MB limit."
        )


def extract_text(filename: str, content: bytes) -> str:
    ext = _ext(filename)
    try:
        if ext == ".pdf":
            return _from_pdf(content)
        if ext in (".docx", ".doc"):
            return _from_docx(content)
        return content.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error("Failed to extract text from %s: %s", filename, e)
        # Last-ditch: try a raw decode so the candidate isn't silently dropped.
        return content.decode("utf-8", errors="ignore")


def _ext(filename: str) -> str:
    name = (filename or "").lower()
    return name[name.rfind("."):] if "." in name else ""


def _from_pdf(content: bytes) -> str:
    import fitz  # PyMuPDF

    parts = []
    with fitz.open(stream=content, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts).strip()


def _from_docx(content: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
